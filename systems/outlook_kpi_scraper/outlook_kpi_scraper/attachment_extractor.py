"""
Attachment-first KPI extraction pipeline.

For each candidate email:
  1. Download attachments to logs/runs/<run_id>/attachments/<msg_entry_id>/
  2. Parse .csv → .xlsx → .xls → .pdf → .docx in priority order.
  3. Return extracted KPI dict with evidence metadata.

Robustness:
  - Safe filename sanitization (Windows invalid chars, path length cap)
  - Absolute paths and os.makedirs before SaveAsFile
  - .xls routed to xlrd engine; .xlsx to openpyxl
  - PDF: pypdf → pdfminer fallback; missing-library logged once (not spammed)
  - DOCX: table + paragraph extraction via python-docx
  - Post-save verification; detailed per-attachment decision logging
"""

import csv
import io
import logging
import os
import re
import threading
import warnings
from typing import Any

from outlook_kpi_scraper.kpi_labels import match_label, KPI_SYNONYMS
from outlook_kpi_scraper.kpi_extractor import parse_money, parse_percent
from outlook_kpi_scraper.ocr_service import extract_pdf_text_with_fallback, ocr_available
from outlook_kpi_scraper.kpi_suitability import compute_suitability

log = logging.getLogger(__name__)

# Suppress noisy PDF warnings (e.g., "Ignoring wrong pointing object")
warnings.filterwarnings("ignore", message=".*wrong pointing object.*")

# Extensions we attempt to parse for KPIs
KPI_EXTENSIONS = {".xlsx", ".xls", ".csv", ".pdf", ".docx"}
# All extensions we download (for debugging / evidence)
DOWNLOAD_EXTENSIONS = {".xlsx", ".xls", ".csv", ".pdf", ".doc", ".docx", ".txt", ".png", ".jpg"}

# Filename keywords that signal KPI-relevant attachments
FILENAME_KPI_KEYWORDS = {
    "report", "financial", "kpi", "dashboard", "weekly", "monthly",
    "cash", "occupancy", "pipeline", "orders", "closings", "revenue",
    "snapshot", "summary", "daily", "p&l", "balance", "income",
    "statement", "model",
}

# Track PDF library availability (log once, not every call)
_pdf_lib_warned = False

# Maximum time (seconds) to spend parsing a single PDF file
PDF_PARSE_TIMEOUT = 30

# Maximum file size (bytes) for PDF parsing  (skip very large PDFs)
PDF_MAX_SIZE_BYTES = 8 * 1024 * 1024  # 8 MB

# Maximum Windows path length safety margin
_MAX_PATH_LEN = 240


# ------------------------------------------------------------------
# Filename safety
# ------------------------------------------------------------------
_INVALID_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
_MAX_FILENAME_LEN = 120
_TRAILING_DOT_SPACE = re.compile(r'[\s.]+$')


def _sanitize_filename(name: str, index: int = 0) -> str:
    """Remove Windows-invalid chars, trailing dots/spaces, and cap total length.

    If *name* is empty or all-invalid, generate a stable fallback name.
    """
    if not name or not name.strip():
        ext = ""
        name = f"attachment_{index}"
    else:
        name = name.strip()

    stem, ext = os.path.splitext(name)
    # Remove invalid chars from stem
    stem = _INVALID_CHARS.sub("_", stem)
    # Remove trailing dots and spaces from stem (Windows restriction)
    stem = _TRAILING_DOT_SPACE.sub("", stem)
    # Collapse multiple underscores
    stem = re.sub(r"_+", "_", stem).strip("_")

    if not stem:
        stem = f"attachment_{index}"

    # Cap length while keeping extension
    if len(stem) + len(ext) > _MAX_FILENAME_LEN:
        stem = stem[:_MAX_FILENAME_LEN - len(ext)]

    return stem + ext


def filename_has_kpi_signal(name: str) -> bool:
    """Return True if the attachment filename contains KPI-relevant keywords."""
    lower = name.lower()
    return any(kw in lower for kw in FILENAME_KPI_KEYWORDS)


# ------------------------------------------------------------------
# Per-attachment decision log (collects entries for CHIP_REVIEW)
# ------------------------------------------------------------------
_attachment_decisions: list[dict] = []


def get_attachment_decisions() -> list[dict]:
    """Return and clear the attachment decision log."""
    global _attachment_decisions
    decisions = list(_attachment_decisions)
    _attachment_decisions = []
    return decisions


def _log_attachment_decision(
    status: str, saved_path: str, original_name: str,
    size: int = 0, error: str = "", engine: str = "",
    suitability: dict | None = None, used_ocr: bool = False,
):
    """Record one attachment save/parse decision for the review log."""
    entry = {
        "status": status,
        "saved_path": saved_path,
        "original_filename": original_name,
        "size": size,
        "error": error,
        "engine": engine,
    }
    if suitability is not None:
        entry["suitability_tier"] = suitability.get("tier")
        entry["suitability_score"] = suitability.get("score")
        entry["suitability_reasons"] = "; ".join(suitability.get("reasons", []))
    if used_ocr:
        entry["used_ocr"] = True
    _attachment_decisions.append(entry)


# ------------------------------------------------------------------
# Public API
# ------------------------------------------------------------------

def extract_kpis_from_attachments(
    outlook_item,
    entry_id: str,
    attachments_dir: str,
) -> dict | None:
    """Download attachments for *outlook_item* and parse for KPI values.

    Returns a dict like ``{"revenue": 12345.0, "evidence": [...]}`` or
    ``None`` if no KPI values were found.
    """
    att_count = _safe_attachment_count(outlook_item)
    if att_count == 0:
        return None

    # Build absolute path: attachments_dir/<safe_entry_id>/
    msg_dir = os.path.join(os.path.abspath(attachments_dir), _safe_dirname(entry_id))
    os.makedirs(msg_dir, exist_ok=True)

    saved_files: list[dict] = []
    for idx in range(1, att_count + 1):  # COM is 1-based
        raw_name = f"attachment_{idx}"
        try:
            att = outlook_item.Attachments.Item(idx)
            raw_name = getattr(att, "FileName", f"attachment_{idx}") or f"attachment_{idx}"
            name = _sanitize_filename(raw_name, index=idx)
            ext = os.path.splitext(name)[1].lower()
            if ext not in DOWNLOAD_EXTENSIONS:
                log.debug("Skipping attachment %s (ext=%s)", name, ext)
                continue

            dest = os.path.abspath(os.path.join(msg_dir, name))

            # Enforce max path length for Windows
            if len(dest) > _MAX_PATH_LEN:
                stem, fext = os.path.splitext(name)
                max_stem = _MAX_PATH_LEN - len(os.path.abspath(msg_dir)) - len(fext) - 2
                if max_stem < 10:
                    max_stem = 10
                name = stem[:max_stem] + fext
                dest = os.path.abspath(os.path.join(msg_dir, name))

            # Ensure parent dir exists (handles any nested edge case)
            os.makedirs(os.path.dirname(dest), exist_ok=True)

            att.SaveAsFile(dest)

            # Post-save verification
            if not os.path.exists(dest):
                error_msg = f"SaveAsFile returned but file does not exist: {dest}"
                log.warning(error_msg)
                _log_attachment_decision("FAILED", dest, raw_name, error=error_msg)
                continue

            file_size = os.path.getsize(dest)
            saved_files.append({"path": dest, "name": name, "ext": ext, "size": file_size})
            _log_attachment_decision("OK", dest, raw_name, size=file_size)
            log.debug("Saved attachment: %s (%d bytes)", dest, file_size)

        except Exception as exc:
            error_msg = str(exc)
            log.warning("Failed to save attachment idx=%d name=%s: %s", idx, raw_name, exc)
            _log_attachment_decision("FAILED", "", raw_name, error=error_msg)

    if not saved_files:
        return None

    # Parse in priority order: csv, xlsx, xls, pdf, docx
    # WITH suitability gating + OCR fallback for scanned PDFs
    kpi: dict[str, Any] = {}
    evidence: list[str] = []

    for item in sorted(saved_files, key=lambda f: _EXT_PRIORITY.get(f["ext"], 99)):
        ext = item["ext"]
        fname = item["name"]
        fpath = item["path"]
        try:
            if ext == ".csv":
                # CSV: quick suitability check on first 5KB
                sample_text = _read_text_sample(fpath, max_bytes=5000)
                suit = compute_suitability(sample_text, filename=fname)
                if suit["tier"] == 4:
                    log.info("SUITABILITY REJECT (Tier 4) csv %s: %s",
                             fname, "; ".join(suit["reasons"]))
                    _log_attachment_decision(
                        "SUIT_REJECT", fpath, fname,
                        size=item.get("size", 0), suitability=suit)
                    continue
                evidence.append(f"suitability csv:{fname} tier={suit['tier']} score={suit['score']}")
                _parse_csv(fpath, kpi, evidence)

            elif ext == ".xlsx":
                suit, sheetnames = _suitability_check_xlsx(fpath, fname)
                if suit["tier"] == 4:
                    log.info("SUITABILITY REJECT (Tier 4) xlsx %s: %s",
                             fname, "; ".join(suit["reasons"]))
                    _log_attachment_decision(
                        "SUIT_REJECT", fpath, fname,
                        size=item.get("size", 0), suitability=suit)
                    continue
                evidence.append(f"suitability xlsx:{fname} tier={suit['tier']} score={suit['score']}")
                _parse_xlsx(fpath, kpi, evidence, preferred_sheets=sheetnames)

            elif ext == ".xls":
                suit_xls, sheetnames_xls = _suitability_check_xls(fpath, fname)
                if suit_xls["tier"] == 4:
                    log.info("SUITABILITY REJECT (Tier 4) xls %s: %s",
                             fname, "; ".join(suit_xls["reasons"]))
                    _log_attachment_decision(
                        "SUIT_REJECT", fpath, fname,
                        size=item.get("size", 0), suitability=suit_xls)
                    continue
                evidence.append(f"suitability xls:{fname} tier={suit_xls['tier']} score={suit_xls['score']}")
                _parse_xls(fpath, kpi, evidence)

            elif ext == ".pdf":
                _parse_pdf_with_suitability(fpath, fname, item.get("size", 0), kpi, evidence)

            elif ext == ".docx":
                _parse_docx(fpath, kpi, evidence)

        except Exception as exc:
            log.warning("Error parsing %s: %s", fname, exc)
            _log_attachment_decision(
                "PARSE_FAILED", fpath, fname,
                size=item.get("size", 0), error=str(exc),
            )

    if not _has_kpi_value(kpi):
        return None

    kpi["evidence"] = evidence
    kpi["attachment_names"] = ";".join(f["name"] for f in saved_files)
    return kpi


def get_attachment_metadata(outlook_item) -> list[dict]:
    """Return metadata for all attachments without downloading."""
    result = []
    att_count = _safe_attachment_count(outlook_item)
    for idx in range(1, att_count + 1):
        try:
            att = outlook_item.Attachments.Item(idx)
            name = getattr(att, "FileName", f"attachment_{idx}")
            ext = os.path.splitext(name)[1].lower()
            size = getattr(att, "Size", 0)
            result.append({"name": name, "ext": ext, "size": size})
        except Exception as exc:
            log.debug("Error reading attachment metadata idx=%d: %s", idx, exc)
    return result


def has_kpi_attachments(outlook_item) -> bool:
    """Return True if item has at least one KPI-relevant attachment."""
    for meta in get_attachment_metadata(outlook_item):
        if meta["ext"] in KPI_EXTENSIONS:
            return True
    return False


# ------------------------------------------------------------------
# Internal: suitability helpers
# ------------------------------------------------------------------

def _read_text_sample(path: str, max_bytes: int = 5000) -> str:
    """Read the first *max_bytes* of a text file for suitability sampling."""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read(max_bytes)
    except Exception:
        return ""


def _suitability_check_xlsx(path: str, filename: str) -> tuple:
    """Quick suitability check for XLSX: read sheet names + sample text.

    Returns (suitability_dict, preferred_sheetnames_list).
    preferred_sheetnames_list contains Summary/Dashboard/KPI/MTD/Report/Census
    sheets that should be parsed first.
    """
    sheetnames: list[str] = []
    sample_text = ""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sheetnames = wb.sheetnames

        # Sample text from first sheet (first 50 rows)
        parts = []
        for ws in wb.worksheets[:2]:
            for row_num, row in enumerate(ws.iter_rows(values_only=True), 1):
                if row_num > 50:
                    break
                parts.append(" ".join(str(c) for c in row if c is not None))
        sample_text = "\n".join(parts)
        wb.close()
    except Exception as exc:
        log.debug("XLSX suitability sampling failed for %s: %s", filename, exc)

    suit = compute_suitability(
        sample_text, filename=filename, sheetnames=sheetnames, is_pdf=False,
    )

    # Determine preferred sheets (Summary/Dashboard/KPI/MTD/Report/Census)
    _PREFERRED = {"summary", "dashboard", "kpi", "mtd", "report", "census"}
    preferred = [s for s in sheetnames if any(p in s.lower() for p in _PREFERRED)]

    return suit, preferred


def _suitability_check_xls(path: str, filename: str) -> tuple:
    """Quick suitability check for XLS: read sheet names + sample text."""
    sheetnames: list[str] = []
    sample_text = ""
    try:
        import xlrd
        wb = xlrd.open_workbook(path)
        sheetnames = wb.sheet_names()

        # Sample first sheet, first 50 rows
        parts = []
        if wb.nsheets > 0:
            ws = wb.sheet_by_index(0)
            for row_num in range(min(ws.nrows, 50)):
                parts.append(" ".join(
                    str(ws.cell_value(row_num, col))
                    for col in range(ws.ncols)
                    if ws.cell_value(row_num, col)
                ))
        sample_text = "\n".join(parts)
    except Exception as exc:
        log.debug("XLS suitability sampling failed for %s: %s", filename, exc)

    suit = compute_suitability(
        sample_text, filename=filename, sheetnames=sheetnames, is_pdf=False,
    )
    return suit, sheetnames


def _parse_pdf_with_suitability(
    path: str, filename: str, file_size: int,
    kpi: dict, evidence: list,
):
    """Parse a PDF with suitability gating and OCR fallback.

    Flow:
      1. Skip if too large / encrypted
      2. Extract text (normal)
      3. Run suitability on extracted text
      4. If Tier 4 → reject
      5. If Tier 3 (OCR candidate) → OCR → re-score
      6. If accepted → scan for KPIs
    """
    # Size check
    if file_size > PDF_MAX_SIZE_BYTES:
        log.info("Skipping large PDF (%d bytes > %d limit): %s",
                 file_size, PDF_MAX_SIZE_BYTES, filename)
        _log_attachment_decision("SKIPPED", path, filename,
                                 size=file_size, error="exceeds size limit")
        return

    # Step 1: normal text extraction
    text, used_ocr = extract_pdf_text_with_fallback(path)
    text_is_empty = len(text.strip()) < 200

    # Step 2: suitability check
    suit = compute_suitability(
        text, filename=filename, is_pdf=True, text_is_empty=text_is_empty,
    )

    # Step 3: Tier 3 → try OCR then re-score
    if suit["tier"] == 3 and suit["used_ocr_candidate_bool"] and not used_ocr:
        log.info("Tier 3 OCR candidate: %s – attempting OCR...", filename)
        if ocr_available():
            from outlook_kpi_scraper.ocr_service import ocr_pdf_first_pages
            ocr_text = ocr_pdf_first_pages(path)
            if ocr_text.strip():
                text = ocr_text
                used_ocr = True
                # Re-score with OCR text
                suit = compute_suitability(
                    text, filename=filename, is_pdf=True, text_is_empty=False,
                )
                log.info("Post-OCR re-score: %s tier=%d score=%d accept=%s",
                         filename, suit["tier"], suit["score"], suit["accept_bool"])
            else:
                log.info("OCR produced no text for %s", filename)
        else:
            log.info("OCR deps not available – cannot OCR Tier 3 candidate: %s", filename)

    # Step 4: Tier 4 → reject
    if suit["tier"] == 4:
        log.info("SUITABILITY REJECT (Tier 4) pdf %s: %s",
                 filename, "; ".join(suit["reasons"]))
        _log_attachment_decision(
            "SUIT_REJECT", path, filename,
            size=file_size, suitability=suit, used_ocr=used_ocr)
        return

    if not suit["accept_bool"] and suit["tier"] != 3:
        log.info("PDF suitability not accepted: %s tier=%d score=%d",
                 filename, suit["tier"], suit["score"])
        _log_attachment_decision(
            "SUIT_SKIP", path, filename,
            size=file_size, suitability=suit, used_ocr=used_ocr)
        return

    evidence.append(
        f"suitability pdf:{filename} tier={suit['tier']} score={suit['score']} ocr={used_ocr}"
    )
    _log_attachment_decision(
        "PARSED", path, filename,
        size=file_size, suitability=suit, used_ocr=used_ocr,
        engine=f"{'ocr' if used_ocr else 'text'}"
    )

    # Step 5: scan for KPIs
    if not text.strip():
        return
    for line_num, line in enumerate(text.splitlines(), 1):
        parts = re.split(r"[:\t|]+|\s{2,}", line)
        _scan_row(parts, kpi, evidence,
                  source=f"pdf:{filename}:line{line_num}")


# ------------------------------------------------------------------
# Internal: parsers
# ------------------------------------------------------------------
_EXT_PRIORITY = {".csv": 1, ".xlsx": 2, ".xls": 3, ".pdf": 4, ".docx": 5}

_MONEY_RE = re.compile(r"[\$]?\s*[\-\(]?\s*[\d,]+\.?\d*\s*[kKmMbB]?\s*\)?")
_NUMBER_RE = re.compile(r"[\-\(]?\s*[\d,]+\.?\d*\s*%?\s*\)?")


def _parse_csv(path: str, kpi: dict, evidence: list):
    """Parse a CSV file for KPI labels + values."""
    log.debug("Parsing CSV: %s", path)
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row_num, row in enumerate(reader, 1):
            _scan_row(row, kpi, evidence, source=f"csv:{os.path.basename(path)}:row{row_num}")


def _parse_xlsx(path: str, kpi: dict, evidence: list, preferred_sheets: list[str] | None = None):
    """Parse an XLSX file for KPI labels + values (openpyxl engine).

    If *preferred_sheets* is provided (e.g. Summary, Dashboard, KPI, MTD),
    those sheets are parsed first to prioritise high-signal tabs.
    """
    log.debug("Parsing XLSX: %s", path)
    try:
        import openpyxl
    except ImportError:
        log.warning("openpyxl not installed – skipping XLSX parsing")
        _log_attachment_decision("PARSE_FAILED", path, os.path.basename(path),
                                 error="openpyxl not installed", engine="none")
        return
    engine_name = f"openpyxl v{getattr(openpyxl, '__version__', 'unknown')}"
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as exc:
        error_msg = f"XLSX parse failed ({engine_name}): {exc}"
        log.warning("%s", error_msg)
        _log_attachment_decision("PARSE_FAILED", path, os.path.basename(path),
                                 error=error_msg, engine=engine_name)
        return

    log.info("XLSX parsed with engine=%s: %s (%d sheets: %s)",
             engine_name, os.path.basename(path), len(wb.sheetnames),
             ", ".join(wb.sheetnames))
    _log_attachment_decision("PARSED", path, os.path.basename(path), engine=engine_name)

    # Order worksheets: preferred sheets first, then the rest
    ordered_ws = list(wb.worksheets)
    if preferred_sheets:
        preferred_lower = {s.lower() for s in preferred_sheets}
        pref = [ws for ws in ordered_ws if ws.title.lower() in preferred_lower]
        rest = [ws for ws in ordered_ws if ws.title.lower() not in preferred_lower]
        ordered_ws = pref + rest
        if pref:
            log.info("XLSX preferred sheets (parsed first): %s",
                     ", ".join(ws.title for ws in pref))

    for ws in ordered_ws:
        for row_num, row in enumerate(ws.iter_rows(values_only=True), 1):
            str_row = [str(c) if c is not None else "" for c in row]
            _scan_row(str_row, kpi, evidence,
                      source=f"xlsx:{os.path.basename(path)}:{ws.title}:row{row_num}")
    wb.close()


def _parse_xls(path: str, kpi: dict, evidence: list):
    """Parse an old-format .xls file for KPI labels + values (xlrd engine).

    Falls back gracefully if xlrd can't parse (e.g., HTML-as-xls, corrupt).
    """
    log.debug("Parsing XLS: %s", path)
    try:
        import xlrd
    except ImportError:
        log.warning("xlrd not installed – skipping .xls parsing. To fix: pip install xlrd")
        _log_attachment_decision("PARSE_FAILED", path, os.path.basename(path),
                                 error="xlrd not installed", engine="none")
        return

    engine_name = f"xlrd v{getattr(xlrd, '__version__', 'unknown')}"
    try:
        wb = xlrd.open_workbook(path)
    except xlrd.biffh.XLRDError as exc:
        # Often happens when .xls is actually HTML or corrupted
        error_msg = f"XLS parse failed ({engine_name}): {exc}"
        log.warning("%s – file saved at %s for inspection", error_msg, path)
        _log_attachment_decision("PARSE_FAILED", path, os.path.basename(path),
                                 error=error_msg, engine=engine_name)
        return
    except Exception as exc:
        error_msg = f"XLS parse failed ({engine_name}): {exc}"
        log.warning("%s", error_msg)
        _log_attachment_decision("PARSE_FAILED", path, os.path.basename(path),
                                 error=error_msg, engine=engine_name)
        return

    log.info("XLS parsed with engine=%s: %s (%d sheets)",
             engine_name, os.path.basename(path), wb.nsheets)
    _log_attachment_decision("PARSED", path, os.path.basename(path), engine=engine_name)

    for ws in wb.sheets():
        for row_num in range(ws.nrows):
            str_row = [str(ws.cell_value(row_num, col)) if ws.cell_value(row_num, col) else ""
                       for col in range(ws.ncols)]
            _scan_row(str_row, kpi, evidence,
                      source=f"xls:{os.path.basename(path)}:{ws.name}:row{row_num + 1}")


def _parse_pdf(path: str, kpi: dict, evidence: list):
    """Extract text from a PDF and scan for KPI labels (no OCR).

    Uses pypdf first, then pdfminer as fallback. Logs missing-library
    warning only once per process to avoid log spam.
    Skips files > PDF_MAX_SIZE_BYTES and enforces PDF_PARSE_TIMEOUT.
    Skips encrypted PDFs.
    """
    global _pdf_lib_warned
    log.debug("Parsing PDF: %s", path)

    # Skip very large PDFs entirely
    try:
        file_size = os.path.getsize(path)
        if file_size > PDF_MAX_SIZE_BYTES:
            log.info("Skipping large PDF (%d bytes > %d limit): %s",
                     file_size, PDF_MAX_SIZE_BYTES, os.path.basename(path))
            evidence.append(f"pdf:{os.path.basename(path)} SKIPPED (too large: {file_size} bytes)")
            _log_attachment_decision("SKIPPED", path, os.path.basename(path),
                                     size=file_size, error="exceeds size limit")
            return
    except OSError:
        pass

    text = ""
    engine_used = ""

    # Try pypdf first (usually fast)
    try:
        import pypdf
        engine_used = f"pypdf v{getattr(pypdf, '__version__', 'unknown')}"
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            reader = pypdf.PdfReader(path)
            # Skip encrypted PDFs
            if reader.is_encrypted:
                log.info("Skipping encrypted PDF: %s", os.path.basename(path))
                evidence.append(f"pdf:{os.path.basename(path)} SKIPPED (encrypted)")
                _log_attachment_decision("SKIPPED", path, os.path.basename(path),
                                         error="encrypted PDF", engine=engine_used)
                return
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
    except ImportError:
        pass
    except Exception as exc:
        log.debug("pypdf failed on %s: %s", os.path.basename(path), exc)

    # Fallback to pdfminer with timeout
    if not text.strip():
        try:
            from pdfminer.high_level import extract_text as pdfm_extract
            import pdfminer
            engine_used = f"pdfminer.six v{getattr(pdfminer, '__version__', 'unknown')}"
            result_container = [None]
            error_container = [None]

            def _do_extract():
                try:
                    result_container[0] = pdfm_extract(path)
                except Exception as e:
                    error_container[0] = e

            t = threading.Thread(target=_do_extract, daemon=True)
            t.start()
            t.join(timeout=PDF_PARSE_TIMEOUT)

            if t.is_alive():
                log.warning("PDF parsing timed out after %ds: %s",
                            PDF_PARSE_TIMEOUT, os.path.basename(path))
                evidence.append(f"pdf:{os.path.basename(path)} TIMEOUT after {PDF_PARSE_TIMEOUT}s")
                _log_attachment_decision("PARSE_FAILED", path, os.path.basename(path),
                                         error=f"timeout after {PDF_PARSE_TIMEOUT}s", engine=engine_used)
                return  # thread is daemon, will be cleaned up on exit

            if error_container[0]:
                log.debug("pdfminer failed on %s: %s", os.path.basename(path), error_container[0])
                _log_attachment_decision("PARSE_FAILED", path, os.path.basename(path),
                                         error=str(error_container[0]), engine=engine_used)
                return

            text = result_container[0] or ""

        except ImportError:
            if not _pdf_lib_warned:
                log.warning("PDF parsing: DISABLED (missing pypdf/pdfminer). To enable: pip install -r requirements.txt")
                _pdf_lib_warned = True
            _log_attachment_decision("SKIPPED", path, os.path.basename(path),
                                     error="DEP_MISSING(PDF)", engine="none")
            return
        except Exception as exc:
            log.debug("pdfminer failed on %s: %s", os.path.basename(path), exc)
            _log_attachment_decision("PARSE_FAILED", path, os.path.basename(path),
                                     error=str(exc), engine=engine_used)
            return

    if not text.strip():
        _log_attachment_decision("PARSE_FAILED", path, os.path.basename(path),
                                 error="no text extracted", engine=engine_used)
        return

    log.info("PDF parsed with engine=%s: %s (%d chars)", engine_used, os.path.basename(path), len(text))
    _log_attachment_decision("PARSED", path, os.path.basename(path), engine=engine_used)

    # Scan line-by-line
    for line_num, line in enumerate(text.splitlines(), 1):
        parts = re.split(r"[:\t|]+|\s{2,}", line)
        _scan_row(parts, kpi, evidence,
                  source=f"pdf:{os.path.basename(path)}:line{line_num}")


def _parse_docx(path: str, kpi: dict, evidence: list):
    """Parse a DOCX file for KPI labels + values (tables + paragraphs)."""
    log.debug("Parsing DOCX: %s", path)
    try:
        from docx import Document
    except ImportError:
        log.warning("python-docx not installed – skipping DOCX parsing")
        return
    doc = Document(path)
    # Tables first (most structured)
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            str_row = [cell.text.strip() for cell in row.cells]
            _scan_row(str_row, kpi, evidence,
                      source=f"docx:{os.path.basename(path)}:table{t_idx + 1}:row{r_idx + 1}")
    # Then paragraphs
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        parts = re.split(r"[:\t|]+|\s{2,}", text)
        _scan_row(parts, kpi, evidence,
                  source=f"docx:{os.path.basename(path)}:para{p_idx + 1}")


# ------------------------------------------------------------------
# Internal: row scanning
# ------------------------------------------------------------------

def _scan_row(cells: list[str], kpi: dict, evidence: list, source: str):
    """Walk cells left-to-right; when a label matches a KPI field, look for
    the nearest numeric value to the right (or in the same cell after a
    separator like ':' or '=' or tab) and store it.

    Fallback: if the label and value share a cell (e.g. "Cash 68,341" from
    PDF/text with no delimiter), extract the first number that appears after
    the matched synonym text.
    """
    for i, cell in enumerate(cells):
        cell_stripped = cell.strip()
        if not cell_stripped:
            continue
        field = match_label(cell_stripped)
        if field is None:
            # Maybe label and value are in the same cell: "Revenue: $1,234"
            parts = re.split(r"[:\-=]\s*", cell_stripped, maxsplit=1)
            if len(parts) == 2:
                field = match_label(parts[0])
                if field and field not in kpi:
                    val = _parse_value(parts[1], field)
                    if val is not None:
                        kpi[field] = val
                        evidence.append(f"{source} cell[{i}] '{cell_stripped}' -> {field}={val}")
                        log.debug("KPI hit: %s=%s from %s", field, val, source)
            continue
        # Label found – look right for numeric value
        if field in kpi:
            continue  # already have this field
        found = False
        for j in range(i + 1, min(i + 4, len(cells))):
            val = _parse_value(cells[j].strip(), field)
            if val is not None:
                kpi[field] = val
                evidence.append(f"{source} cell[{i}]->cell[{j}] '{cell_stripped}'->'{cells[j].strip()}' -> {field}={val}")
                log.debug("KPI hit: %s=%s from %s", field, val, source)
                found = True
                break
        # Fallback: label and value share the same cell (PDF financial layouts)
        if not found:
            val = _extract_value_after_label(cell_stripped, field)
            if val is not None:
                kpi[field] = val
                evidence.append(f"{source} cell[{i}] '{cell_stripped}' (same-cell) -> {field}={val}")
                log.debug("KPI hit (same-cell): %s=%s from %s", field, val, source)


def _extract_value_after_label(cell_text: str, field: str):
    """Extract the first numeric value that appears *after* the KPI synonym
    within the same cell text.  Handles PDF layouts like 'Cash 68,341' where
    label and number are separated only by whitespace.
    """
    from outlook_kpi_scraper.kpi_labels import _REVERSE
    cell_lower = cell_text.lower()
    # Try the longest matching synonym first for precision
    for syn in sorted((s for s, f in _REVERSE.items() if f == field), key=len, reverse=True):
        idx = cell_lower.find(syn)
        if idx < 0:
            continue
        after = cell_text[idx + len(syn):]
        # Find the first money-like token in the remainder
        m = _MONEY_RE.search(after)
        if m:
            val = _parse_value(m.group(0).strip(), field)
            if val is not None:
                return val
        break  # only try the best (longest) matching synonym
    return None


def _parse_value(raw: str, field: str):
    """Parse a raw string into the appropriate numeric type for *field*."""
    if not raw:
        return None
    raw = raw.strip()
    if field == "occupancy":
        if "%" in raw:
            return parse_percent(raw)
        v = parse_money(raw)
        if v is not None and 0 < v <= 100:
            return v / 100.0
        return v
    if "count" in field:
        v = parse_money(raw)
        return int(v) if v is not None else None
    return parse_money(raw)


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _has_kpi_value(kpi: dict) -> bool:
    """Return True if at least one canonical KPI field has a value."""
    kpi_fields = set(KPI_SYNONYMS.keys())
    return any(kpi.get(f) is not None for f in kpi_fields)


def _safe_attachment_count(item) -> int:
    try:
        return item.Attachments.Count
    except Exception:
        return 0


def _safe_dirname(entry_id: str) -> str:
    """Create a filesystem-safe directory name from an EntryID."""
    safe = re.sub(r"[^A-Za-z0-9_\-]", "", entry_id or "unknown")
    return safe[-24:] if len(safe) > 24 else safe
